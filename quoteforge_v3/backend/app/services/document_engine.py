"""
Document Rendering Engine — generates PDF and DOCX documents
using ReportLab and python-docx.

The HTML master-template path lives in render_pdf_from_html() below: when
a tenant has an active master template (Template.is_master=True with
html_body), quote generation routes through Jinja2 + xhtml2pdf instead
of the section-based ReportLab pipeline. The legacy render_pdf is kept
as a fallback when no master template is configured.
"""
import io
import logging
from pathlib import Path
from datetime import datetime

logger = logging.getLogger(__name__)

GENERATED_DIR = Path(__file__).resolve().parent.parent.parent / "generated_docs"
GENERATED_DIR.mkdir(exist_ok=True)


def render_pdf_from_html(doc_id: str, html_template: str, context: dict) -> str:
    """Render a PDF from an HTML/Jinja2 master template.

    `context` should contain the same keys the default master template
    references (client_name, deal_name, line_items, subtotal, discount,
    tax, total, currency_symbol, valid_until, etc.) — extras are simply
    ignored by Jinja2's strict-undefined-off behaviour.

    Uses xhtml2pdf (PISA) so there are no system dependencies on
    Cairo/wkhtmltopdf — works unchanged in the Render Docker image.
    """
    from jinja2 import Environment, BaseLoader, StrictUndefined  # noqa: F401
    from xhtml2pdf import pisa

    # Lenient undefined → missing context keys render as empty string,
    # which matches the "show what you have" semantics admins expect when
    # they tweak the template. Swap to StrictUndefined if you want loud
    # errors on missing variables.
    env = Environment(loader=BaseLoader(), autoescape=False)
    template = env.from_string(html_template)
    rendered_html = template.render(**context)

    filepath = str(GENERATED_DIR / f"{doc_id}.pdf")
    with open(filepath, "wb") as fp:
        pisa_status = pisa.CreatePDF(src=rendered_html, dest=fp)

    if pisa_status.err:
        # PISA reports recoverable rendering errors via .err; surface
        # them so the caller can fall back to the legacy renderer
        # instead of returning a half-baked PDF.
        raise RuntimeError(
            f"xhtml2pdf failed to render master template for {doc_id}: "
            f"{pisa_status.err} errors"
        )
    logger.info("HTML→PDF generated: %s", filepath)
    return filepath


def render_pdf(doc_id: str, sections: dict, pricing: dict, context: dict) -> str:
    """Render a professional PDF proposal document."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable, PageBreak,
    )

    filepath = str(GENERATED_DIR / f"{doc_id}.pdf")
    doc = SimpleDocTemplate(filepath, pagesize=letter,
                            topMargin=0.75*inch, bottomMargin=0.75*inch,
                            leftMargin=inch, rightMargin=inch)

    styles = getSampleStyleSheet()
    brand_color = HexColor("#3576e8")

    # Custom styles
    title_style = ParagraphStyle("Title2", parent=styles["Title"],
                                  textColor=brand_color, fontSize=24, spaceAfter=6)
    heading_style = ParagraphStyle("Heading2Custom", parent=styles["Heading2"],
                                    textColor=brand_color, fontSize=16, spaceBefore=20, spaceAfter=10)
    body_style = styles["BodyText"]
    body_style.fontSize = 11
    body_style.leading = 16

    elements = []

    # Title page
    elements.append(Spacer(1, 2*inch))
    elements.append(Paragraph("QuoteForge", title_style))
    elements.append(Spacer(1, 12))
    doc_type = context.get("type", "Proposal")
    elements.append(Paragraph(f"<b>{doc_type}</b> — {context.get('deal_name', 'Project')}", styles["Heading3"]))
    elements.append(Spacer(1, 24))
    elements.append(Paragraph(f"Prepared for: <b>{context.get('client_name', 'Client')}</b>", body_style))
    elements.append(Paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", body_style))
    elements.append(Paragraph(f"Document ID: {doc_id}", body_style))
    # Validity
    valid_until = context.get('valid_until')
    if valid_until:
        from datetime import datetime as _dt
        if isinstance(valid_until, str):
            try:
                valid_until = _dt.fromisoformat(valid_until.replace('Z', '+00:00'))
            except Exception:
                valid_until = None
        if valid_until:
            elements.append(Paragraph(
                f"<b>Valid Until:</b> <font color='#d97706'>{valid_until.strftime('%B %d, %Y')}</font> (30 days)",
                body_style
            ))
    elements.append(Spacer(1, 12))
    elements.append(HRFlowable(width="100%", color=brand_color, thickness=2))
    elements.append(PageBreak())

    # Sections
    section_order = ["Cover Letter", "Summary", "Scope", "Deliverables", "Pricing", "Terms"]
    for section_name in section_order:
        content = sections.get(section_name, "")
        if not content:
            continue
        elements.append(Paragraph(section_name, heading_style))
        elements.append(HRFlowable(width="100%", color=HexColor("#e5e7eb"), thickness=0.5))
        elements.append(Spacer(1, 8))
        for para in content.split("\n"):
            para = para.strip()
            if para:
                # Escape special characters for ReportLab
                para = para.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                elements.append(Paragraph(para, body_style))
                elements.append(Spacer(1, 4))
        elements.append(Spacer(1, 12))

    # Pricing table
    if pricing.get("subtotal"):
        elements.append(Paragraph("Pricing Breakdown", heading_style))
        elements.append(HRFlowable(width="100%", color=HexColor("#e5e7eb"), thickness=0.5))
        elements.append(Spacer(1, 8))

        table_data = [["Item", "Amount"]]

        # Line items
        for item in context.get("line_items", []):
            if isinstance(item, dict):
                name = item.get("product", "Item")
                qty = item.get("quantity", 1)
                price = item.get("unit_price", 0)
                table_data.append([f"{name} x {qty}", f"${price * qty:,.2f}"])

        table_data.append(["Subtotal", f"${pricing['subtotal']:,.2f}"])
        if pricing.get("discount", 0) > 0:
            table_data.append(["Discount", f"-${pricing['discount']:,.2f}"])
        if pricing.get("tax", 0) > 0:
            table_data.append(["Tax", f"${pricing['tax']:,.2f}"])
        table_data.append(["Total", f"${pricing['total']:,.2f}"])

        t = Table(table_data, colWidths=[4*inch, 2*inch])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), brand_color),
            ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#ffffff")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("ALIGN", (1, 0), (1, -1), "RIGHT"),
            ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#e5e7eb")),
            ("BACKGROUND", (0, -1), (-1, -1), HexColor("#f3f4f6")),
            ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
            ("TOPPADDING", (0, 0), (-1, -1), 8),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ]))
        elements.append(t)

    # Footer
    elements.append(Spacer(1, 40))
    elements.append(HRFlowable(width="100%", color=brand_color, thickness=1))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph(
        f"Generated by QuoteForge | {datetime.now().strftime('%B %d, %Y %I:%M %p')}",
        ParagraphStyle("Footer", parent=body_style, fontSize=8, textColor=HexColor("#9ca3af")),
    ))

    doc.build(elements)
    logger.info(f"PDF generated: {filepath}")
    return filepath


def render_docx(doc_id: str, sections: dict, pricing: dict, context: dict) -> str:
    """Render a professional DOCX proposal document."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    filepath = str(GENERATED_DIR / f"{doc_id}.docx")
    doc = Document()

    # Style setup
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    brand = RGBColor(0x35, 0x76, 0xe8)

    # Title
    title = doc.add_heading("QuoteForge", level=0)
    title.runs[0].font.color.rgb = brand

    doc_type = context.get("type", "Proposal")
    sub = doc.add_paragraph()
    run = sub.add_run(f"{doc_type} — {context.get('deal_name', 'Project')}")
    run.font.size = Pt(14)
    run.font.color.rgb = brand

    doc.add_paragraph(f"Prepared for: {context.get('client_name', 'Client')}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph(f"Document ID: {doc_id}")

    # Validity
    valid_until = context.get('valid_until')
    if valid_until:
        from datetime import datetime as _dt
        if isinstance(valid_until, str):
            try:
                valid_until = _dt.fromisoformat(valid_until.replace('Z', '+00:00'))
            except Exception:
                valid_until = None
        if valid_until:
            p = doc.add_paragraph()
            p.add_run("Valid Until: ").bold = True
            run = p.add_run(f"{valid_until.strftime('%B %d, %Y')} (30 days)")
            run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
    doc.add_page_break()

    # Sections
    section_order = ["Cover Letter", "Summary", "Scope", "Deliverables", "Pricing", "Terms"]
    for section_name in section_order:
        content = sections.get(section_name, "")
        if not content:
            continue
        heading = doc.add_heading(section_name, level=1)
        heading.runs[0].font.color.rgb = brand
        for para in content.split("\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)

    # Pricing table
    if pricing.get("subtotal"):
        heading = doc.add_heading("Pricing Breakdown", level=1)
        heading.runs[0].font.color.rgb = brand

        rows = [["Item", "Amount"]]
        for item in context.get("line_items", []):
            if isinstance(item, dict):
                name = item.get("product", "Item")
                qty = item.get("quantity", 1)
                price = item.get("unit_price", 0)
                rows.append([f"{name} x {qty}", f"${price * qty:,.2f}"])
        rows.append(["Subtotal", f"${pricing['subtotal']:,.2f}"])
        if pricing.get("discount", 0) > 0:
            rows.append(["Discount", f"-${pricing['discount']:,.2f}"])
        if pricing.get("tax", 0) > 0:
            rows.append(["Tax", f"${pricing['tax']:,.2f}"])
        rows.append(["Total", f"${pricing['total']:,.2f}"])

        table = doc.add_table(rows=len(rows), cols=2)
        table.style = "Light Grid Accent 1"
        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                table.rows[i].cells[j].text = cell_text

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    run = footer.add_run(f"Generated by QuoteForge | {datetime.now().strftime('%B %d, %Y %I:%M %p')}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    doc.save(filepath)
    logger.info(f"DOCX generated: {filepath}")
    return filepath
